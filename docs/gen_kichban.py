import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Colour palette ───────────────────────────────────────────────────────────
GREEN_DARK  = RGBColor(0x14, 0x63, 0x2D)   # VKU dark green
GREEN_MID   = RGBColor(0x1D, 0x7A, 0x3A)
GREEN_LIGHT = RGBColor(0xE8, 0xF5, 0xE9)
TEAL        = RGBColor(0x00, 0x7B, 0x83)
ORANGE      = RGBColor(0xE6, 0x5C, 0x00)
GRAY_DARK   = RGBColor(0x37, 0x37, 0x37)
GRAY_MID    = RGBColor(0x75, 0x75, 0x75)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helper: set paragraph shading ────────────────────────────────────────────
def shade_paragraph(para, hex_color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

# ── Helper: add a run with full formatting ────────────────────────────────────
def add_run(para, text, bold=False, italic=False, color=None, size=11, font='Calibri'):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

# ═══════════════════════════════════════════════════════════════
# COVER / TITLE PAGE
# ═══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(p, '146320')
add_run(p, '\n📋  KỊCH BẢN THUYẾT TRÌNH  📋\n', bold=True, color=WHITE, size=22, font='Calibri')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(p, '1D7A3A')
add_run(p, 'ĐỒ ÁN CHUYÊN NGÀNH 3 — VKU MARKET\n', bold=True, color=WHITE, size=15, font='Calibri')
add_run(p, 'Ứng dụng Blockchain trong quản lý & truy xuất nguồn gốc chuỗi cung ứng nông sản\n', italic=True, color=RGBColor(0xC8,0xE6,0xC9), size=12, font='Calibri')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(p, 'E8F5E9')
add_run(p, '👤  Lê Nhật Huy  (23IT102)     |     Nguyễn Thị Thùy Tiến  (23IT273)\n', bold=True, color=GREEN_DARK, size=12, font='Calibri')
add_run(p, 'Lớp IT23A  ·  Trường Đại học Công nghệ Thông tin và Truyền thông Việt – Hàn (VKU)', color=GRAY_DARK, size=11, font='Calibri')

# Legend
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '\n🔵 [HUY] = Huy nói     |     🟣 [TIẾN] = Tiến nói     |     ▶ [CHUYỂN SLIDE] = click sang slide tiếp\n', italic=True, color=GRAY_MID, size=10)

doc.add_paragraph()  # spacer

# ═══════════════════════════════════════════════════════════════
# DATA — each slide is a dict
# ═══════════════════════════════════════════════════════════════
slides = [
    {
        "number": 1,
        "title": "Trang bìa / Cover",
        "lines": [
            ("HUY", "Kính chào quý thầy cô trong hội đồng chấm điểm. Em xin tự giới thiệu, em là Lê Nhật Huy, mã sinh viên 23IT102."),
            ("TIẾN", "Và em là Nguyễn Thị Thùy Tiến, mã sinh viên 23IT273. Hôm nay hai em rất vinh dự được trình bày đồ án chuyên ngành 3 với chủ đề: 'VKU Market — Xây dựng hệ thống quản lý và truy xuất nguồn gốc chuỗi cung ứng nông sản ứng dụng công nghệ Blockchain.' Kính mời thầy cô theo dõi phần thuyết trình của hai em."),
        ],
        "note": ""
    },
    {
        "number": 2,
        "title": "Mục lục / Nội dung trình bày",
        "lines": [
            ("HUY", "Thưa thầy cô, phần trình bày hôm nay của nhóm em sẽ được cấu trúc theo 5 phần chính:\n  • Phần 1 — Đặt vấn đề: Lý do thực tế khiến nhóm em chọn đề tài này.\n  • Phần 2 — Giải pháp và Kiến trúc hệ thống.\n  • Phần 3 — Các tính năng cốt lõi đã hoàn thiện.\n  • Phần 4 — Kết quả đạt được và demo sản phẩm.\n  • Phần 5 — Hướng phát triển tương lai.\nSau mỗi phần, thầy cô có thể đặt câu hỏi thêm. Nhóm em rất sẵn sàng giải đáp. Em xin phép bắt đầu."),
        ],
        "note": ""
    },
    {
        "number": 3,
        "title": "Đặt vấn đề / Thực trạng chuỗi cung ứng",
        "lines": [
            ("TIẾN", "Thưa thầy cô, trong thực tế hiện nay, ngành nông sản Việt Nam đang phải đối mặt với một vòng luẩn quẩn rất nan giải.\n\nNgười tiêu dùng thì lo lắng về nguồn gốc thực phẩm: 'Quả xoài này trồng ở đâu? Có phun thuốc không? Shipper có bảo quản đúng cách không?' Họ không có cách nào kiểm chứng được, vì tất cả thông tin đều nằm trong một chiếc giấy tờ giấy có thể dễ dàng bị làm giả.\n\nVề phía người nông dân và doanh nghiệp, họ đầu tư trồng trọt chất lượng cao nhưng không thể chứng minh được điều đó, dẫn đến bị ép giá hoặc bị mất cạnh tranh với hàng kém chất lượng giá rẻ.\n\nChính vì bài toán 'bất cân xứng thông tin' này mà nhóm em đã nghiên cứu và đề xuất VKU Market làm giải pháp."),
        ],
        "note": ""
    },
    {
        "number": 4,
        "title": "Công nghệ sử dụng / Tech Stack",
        "lines": [
            ("HUY", "Thưa thầy cô, để xây dựng hệ thống này, nhóm em đã lựa chọn một bộ công nghệ hiện đại theo từng tầng rõ ràng.\n\nVề tầng Frontend, nhóm em sử dụng Next.js 14 với kiến trúc App Router làm nền tảng chính. Wagmi và Viem là hai thư viện chuyên dụng để tương tác với ví MetaMask và đọc dữ liệu từ blockchain Ethereum một cách type-safe và an toàn.\n\nVề tầng Backend và Dữ liệu, nhóm em dùng Prisma ORM kết hợp MySQL để lưu trữ dữ liệu nghiệp vụ với tốc độ truy vấn cao. Toàn bộ API được viết theo chuẩn REST trong Next.js API Routes.\n\nVề tầng Blockchain, Smart Contract được viết bằng Solidity và đã được triển khai lên mạng thử nghiệm Ethereum Sepolia.\n\nĐặc biệt, nhóm em còn tích hợp thêm Lớp trí tuệ nhân tạo sử dụng Google Gemini API, với cơ chế dự phòng thông qua OpenRouter để đảm bảo hệ thống AI luôn hoạt động ngay cả khi đường truyền mạng gặp sự cố."),
        ],
        "note": ""
    },
    {
        "number": 5,
        "title": "Phân tách dữ liệu: On-chain, Off-chain và AI ✨ (slide mới bổ sung)",
        "lines": [
            ("TIẾN", "Thưa thầy cô, một trong những quyết định thiết kế quan trọng nhất của nhóm em là việc phân tách dữ liệu thành ba lớp riêng biệt, có mục đích rõ ràng và bổ sung lẫn nhau.\n\nLớp thứ nhất là Off-chain, sử dụng MySQL. Đây là nơi lưu trữ toàn bộ dữ liệu nghiệp vụ như thông tin tài khoản người dùng, danh mục sản phẩm, đơn hàng và tồn kho. Lý do nhóm em không đưa tất cả lên blockchain là vì chi phí gas rất cao và tốc độ truy vấn chậm, không đáp ứng được yêu cầu thời gian thực của người dùng.\n\nLớp thứ hai là On-chain, trên Ethereum Sepolia. Đây là nơi lưu trữ các bằng chứng quan trọng nhất — những thứ cần tính BẤT BIẾN và KHÔNG THỂ GIẢ MẠO: đó là lịch sử hành trình của từng lô hàng, chữ ký số của shipper, và trạng thái giao nhận.\n\nLớp thứ ba là Lớp xử lý AI — đây là điểm độc đáo mà nhóm em tự hào. AI đóng vai trò cầu nối thông minh giữa hai lớp dữ liệu trên, lấy dữ liệu thô và biến thành thông tin có giá trị thực tế cho người dùng."),
        ],
        "note": "💡 Slide này đã bổ sung Lớp AI — nhớ nói rõ đây là tính năng ĐÃ hoàn thiện, không phải tương lai."
    },
    {
        "number": 6,
        "title": "Mô hình Web2.5 Hybrid",
        "lines": [
            ("HUY", "Từ sự phân tách dữ liệu đó, nhóm em xác định mô hình kiến trúc tổng thể là Web2.5 Hybrid, hay còn có thể gọi là 'Tin tưởng nhưng Xác minh' — Trust but Verify.\n\nVề cơ bản, tất cả các thao tác hàng ngày như tìm kiếm sản phẩm, xem giá cả, quản lý đơn hàng đều được xử lý ở tầng Web2 truyền thống với tốc độ phản hồi dưới 50 mili-giây. Người dùng sẽ không bao giờ phải chờ đợi.\n\nNhưng khi cần XÁC MINH một thông tin quan trọng — ví dụ như 'Lô hàng sầu riêng này có thực sự được giao không?' — hệ thống sẽ tra cứu ngay lên blockchain, nơi không ai có thể gian lận. Đây là cơ chế hybrid mà nhóm em thiết kế."),
        ],
        "note": ""
    },
    {
        "number": 7,
        "title": "Kiến trúc App Shell SPA",
        "lines": [
            ("TIẾN", "Thưa thầy cô, đây là một trong những vấn đề kỹ thuật thú vị nhất mà nhóm em gặp phải trong quá trình xây dựng.\n\nVấn đề là: khi người dùng điều hướng giữa các trang trong một ứng dụng Web3 truyền thống kiểu Multi-Page, trình duyệt sẽ tải lại toàn bộ trang — điều này khiến tất cả các Provider kết nối ví MetaMask bị hủy và phải kết nối lại từ đầu. Rất bực bội cho người dùng!\n\nGiải pháp của nhóm em là thiết kế theo kiến trúc App Shell. Nhóm em đặt toàn bộ ứng dụng trong một tệp page.tsx duy nhất, biến nó thành một Single-Page Application. Các Provider kết nối ví — Wagmi, RainbowKit, Theme Provider — được gắn cố định tại Root Layout, và sẽ KHÔNG BAO GIỜ bị hủy trong suốt phiên làm việc của người dùng.\n\nKhi người dùng click chuyển từ 'Cửa hàng' sang 'Quản lý kho', hệ thống chỉ đơn giản là thay đổi một biến trạng thái activeView, không hề tải lại trang. Kết quả là ví MetaMask luôn kết nối liên tục, trải nghiệm mượt mà như một ứng dụng di động."),
        ],
        "note": ""
    },
    {
        "number": 8,
        "title": "Phân quyền RBAC — 4 vai trò",
        "lines": [
            ("HUY", "Thưa thầy cô, hệ thống VKU Market được thiết kế phục vụ 4 nhóm đối tượng người dùng với phân quyền rõ ràng theo mô hình RBAC — Role-Based Access Control.\n\n• ADMIN — người quản trị hệ thống với toàn quyền: duyệt đơn hàng COD, quản lý người dùng và xem báo cáo tổng hợp.\n\n• SUPPLIER — người nông dân hoặc nhà cung cấp. Họ có quyền đăng ký sản phẩm lên hệ thống và KÝ GIAO DỊCH blockchain để tạo lô hàng mới. Mỗi lô hàng được ghi lên Ethereum Sepolia với chữ ký số của người tạo.\n\n• SHIPPER — người vận chuyển. Khi nhận hàng, shipper sẽ dùng ví MetaMask để ký cập nhật trạng thái vận chuyển trực tiếp lên blockchain. Nhờ đó, địa chỉ ví của shipper được ghi lại vĩnh viễn — đây chính là cơ chế xác thực danh tính shipper không thể giả mạo.\n\n• CUSTOMER — khách hàng. Điểm hay ở đây là khách hàng KHÔNG CẦN có ví crypto. Họ có thể mua sắm bình thường như một sàn thương mại điện tử thông thường."),
        ],
        "note": ""
    },
    {
        "number": 9,
        "title": "Luồng hoạt động của Supplier",
        "lines": [
            ("TIẾN", "Thưa thầy cô, để cụ thể hóa hơn, em xin trình bày luồng hoạt động của vai trò Supplier — người nông dân.\n\nĐầu tiên, nhà cung cấp đăng nhập và tạo một sản phẩm mới: điền tên, xuất xứ, giá cả. Thông tin này được lưu vào MySQL Off-chain.\n\nSau đó, khi đã có sản phẩm, nhà cung cấp tạo một lô hàng mới và nhấn nút 'Ký & Tạo lô hàng'. Lúc này, MetaMask sẽ pop-up lên yêu cầu xác nhận. Khi nhà cung cấp xác nhận, Smart Contract trên Sepolia sẽ ghi lại thông tin lô hàng đó vĩnh viễn trên blockchain, kèm theo địa chỉ ví của người tạo làm bằng chứng xuất xứ.\n\nTừ thời điểm này, lô hàng này có một 'hộ chiếu blockchain' — một mã ID duy nhất mà bất kỳ ai cũng có thể dùng để tra cứu lịch sử đầy đủ của nó."),
        ],
        "note": ""
    },
    {
        "number": 10,
        "title": "Luồng hoạt động của Shipper",
        "lines": [
            ("HUY", "Tiếp theo là luồng của Shipper. Đây là phần mà nhóm em thấy có tính ứng dụng thực tế rất cao.\n\nKhi nhận được đơn giao hàng, shipper đăng nhập vào hệ thống với ví MetaMask của cá nhân mình. Sau đó, shipper cập nhật trạng thái lô hàng: từ 'Đang vận chuyển' đến 'Đã giao hàng'. Mỗi lần cập nhật là một lần ký giao dịch blockchain, ghi lại chính xác: vị trí địa lý, thời gian, tình trạng hàng hóa, và địa chỉ ví của người cập nhật.\n\nĐiều này giải quyết một vấn đề thực tế rất lớn: 'Con shipper này có thực sự đã đến địa điểm đó không? Có tự ý thay đổi lộ trình không?' Với blockchain, mọi thứ đều được ghi lại bất biến và công khai, đảm bảo trách nhiệm giải trình của từng cá nhân trong chuỗi."),
        ],
        "note": ""
    },
    {
        "number": 11,
        "title": "Smart Contract — Thiết kế và triển khai",
        "lines": [
            ("TIẾN", "Thưa thầy cô, đây là phần kỹ thuật quan trọng nhất — Smart Contract.\n\nNhóm em đã viết một Smart Contract bằng Solidity với ba hàm chính:\n\n• Hàm createProduct: Nhận vào ID lô hàng, tên sản phẩm và xuất xứ. Ghi lại lịch sử đầu tiên với trạng thái 'Đã tạo' và lưu địa chỉ ví người tạo.\n\n• Hàm updateStatus: Chỉ những địa chỉ ví đã được cấp quyền mới có thể gọi hàm này. Hàm thêm một bản ghi mới vào mảng lịch sử của lô hàng đó, ghi lại trạng thái mới, thời gian, vị trí và người cập nhật.\n\n• Hàm getProduct: Hàm đọc công khai — bất kỳ ai cũng có thể gọi để xem toàn bộ lịch sử của một lô hàng chỉ với ID của nó, mà không cần tốn phí gas.\n\nSmart Contract này đã được deploy lên Ethereum Sepolia Testnet tại địa chỉ cụ thể và có thể tra cứu công khai trên Etherscan."),
        ],
        "note": ""
    },
    {
        "number": 12,
        "title": "Smart Contract — Cơ chế bảo mật",
        "lines": [
            ("HUY", "Thưa thầy cô, về bảo mật, nhóm em đã xây dựng cơ chế phân quyền ngay trong Smart Contract.\n\nTrong Solidity, nhóm em dùng modifier onlyAuthorized để chặn các địa chỉ ví không có quyền. Chỉ những ví đã được ADMIN cấp quyền từ trước mới có thể ký giao dịch cập nhật trạng thái lô hàng.\n\nĐiều này có nghĩa là dù có ai đó biết địa chỉ hợp đồng, họ cũng không thể tự ý thêm hoặc sửa dữ liệu vào lô hàng của người khác. Bảo mật được thực thi ở tầng blockchain, không phải chỉ ở tầng ứng dụng."),
        ],
        "note": ""
    },
    {
        "number": 13,
        "title": "Smart Inventory Automation",
        "lines": [
            ("TIẾN", "Thưa thầy cô, đây là tính năng mà nhóm em đặc biệt tâm đắc vì nó giải quyết một bài toán tự động hóa rất hay.\n\nKịch bản như sau: Shipper giao hàng thành công và ký giao dịch 'Đã giao hàng' lên blockchain. Vậy làm sao để hệ thống tự động biết và cập nhật số lượng tồn kho trong MySQL?\n\nNhóm em đã xây dựng một cơ chế tự động gọi là Smart Inventory Automation. Khi Frontend nhận được Transaction Receipt xác nhận thành công từ blockchain, nó ngay lập tức gọi một API endpoint nội bộ. Backend nhận tín hiệu này, xác minh lại trên blockchain để đảm bảo không bị giả mạo, rồi dùng Prisma Upsert để tự động tăng số lượng trong kho tương ứng.\n\nĐặc biệt, nhóm em còn xây dựng cơ chế chống ghi dồn hai lần — idempotency. Hệ thống kiểm tra xem transaction hash này đã được xử lý chưa trước khi cập nhật. Điều này đảm bảo dù người dùng click nhiều lần, kho hàng chỉ được cộng đúng một lần."),
        ],
        "note": ""
    },
    {
        "number": 14,
        "title": "Luồng đặt hàng và thanh toán",
        "lines": [
            ("HUY", "Thưa thầy cô, về trải nghiệm mua sắm của khách hàng, nhóm em đã thiết kế sao cho đơn giản nhất có thể.\n\nKhách hàng vào trang cửa hàng, chọn sản phẩm, thêm vào giỏ hàng và đặt hàng. Họ có thể lựa chọn thanh toán theo hai hình thức:\n\n• COD — thanh toán khi nhận hàng. Admin sẽ duyệt đơn thủ công.\n• Chuyển khoản thông thường.\n\nSau khi đơn hàng được duyệt, hệ thống sẽ cập nhật trạng thái và shipper được phân công giao hàng. Toàn bộ hành trình này có thể theo dõi theo thời gian thực thông qua trang Truy Xuất Nguồn Gốc, nơi dữ liệu blockchain được hiển thị minh bạch."),
        ],
        "note": ""
    },
    {
        "number": 15,
        "title": "Trải nghiệm người dùng & Tích hợp AI ✨ (đã bổ sung AI)",
        "lines": [
            ("TIẾN", "Thưa thầy cô, về phần trải nghiệm người dùng, nhóm em đã tập trung vào bốn khía cạnh quan trọng.\n\nThứ nhất, về giao diện — VKU Market được thiết kế theo phong cách hiện đại, tối giản, với hiệu ứng chuyển động mượt mà. Mục tiêu là người dùng không cảm thấy mình đang dùng 'phần mềm quản lý kho' mà là đang mua sắm trên một sàn thương mại điện tử chuyên nghiệp.\n\nThứ hai, về hiệu năng — nhóm em đã tối ưu hóa rất kỹ bằng cách áp dụng React.memo và useCallback để tránh re-render thừa, kết hợp với TanStack Query để cache API response. Kết quả là tốc độ cuộn sản phẩm và chuyển đổi giao diện đạt 60 khung hình mỗi giây.\n\nThứ ba, về truy xuất nguồn gốc — người dùng chỉ cần nhập mã lô hàng là xem được toàn bộ hành trình sản phẩm từ vườn đến tay, được đọc từ blockchain theo thời gian thực.\n\nThứ tư — đây là điểm độc đáo nhất — là Trợ lý mua sắm AI cá nhân hóa. Ngay khi vào trang chủ, hệ thống đọc vị trí GPS và thời tiết thực tế của người dùng, rồi AI sẽ tư vấn loại nông sản nào đang vào mùa, tốt nhất cho sức khỏe trong điều kiện thời tiết hiện tại."),
        ],
        "note": "💡 Nhấn mạnh: đây là tính năng ĐÃ HOÀN THIỆN, thực sự chạy được trên web."
    },
    {
        "number": 16,
        "title": "Khó khăn kỹ thuật và cách giải quyết",
        "lines": [
            ("HUY", "Thưa thầy cô, trong quá trình xây dựng, nhóm em gặp phải một số khó khăn kỹ thuật thú vị.\n\nVấn đề đầu tiên là BigInt Serialization. Khi đọc dữ liệu từ Smart Contract, các giá trị ID và timestamp trả về dưới dạng BigInt — một kiểu dữ liệu mà JavaScript's JSON.stringify không thể xử lý và sẽ bị crash. Nhóm em giải quyết bằng cách tạo một tầng chuyển đổi dữ liệu, tự động convert tất cả BigInt thành String trước khi đưa vào React state.\n\nVấn đề thứ hai là sự không đồng bộ giữa blockchain và database. Blockchain cần khoảng 15-30 giây để xác nhận một giao dịch, trong khi người dùng muốn thấy phản hồi ngay lập tức. Nhóm em giải quyết bằng cách sử dụng Optimistic UI — hiển thị kết quả dự kiến ngay, sau đó xác nhận lại khi blockchain phản hồi.\n\nVấn đề thứ ba là bảo mật API. Để ngăn chặn các cuộc tấn công giả mạo blockchain event, tất cả webhook từ Frontend đều phải được Backend xác minh lại trực tiếp với blockchain Sepolia trước khi thực thi bất kỳ thay đổi nào trong database."),
        ],
        "note": ""
    },
    {
        "number": 17,
        "title": "Hướng phát triển tương lai ✨ (đã cập nhật)",
        "lines": [
            ("TIẾN", "Thưa thầy cô, do tính năng AI tư vấn nông sản theo mùa đã được nhóm em hoàn thiện trong phiên bản hiện tại, hướng phát triển tương lai sẽ tập trung vào ba hướng lớn hơn.\n\nHướng thứ nhất là mở rộng lên Layer 2. Hiện tại nhóm em đang chạy trên Sepolia Testnet là mạng thử nghiệm. Khi ra production thực tế, nhóm em dự định chuyển sang Arbitrum hoặc Base — đây là các giải pháp Layer 2 của Ethereum, giảm phí gas khoảng 95% so với Mainnet.\n\nHướng thứ hai là tích hợp IoT Sensor theo dõi chuỗi lạnh. Mỗi xe tải vận chuyển nông sản sẽ được gắn cảm biến đo nhiệt độ và độ ẩm. Dữ liệu từ cảm biến sẽ được ghi trực tiếp lên blockchain theo thời gian thực.\n\nHướng thứ ba là ứng dụng học máy cho Dự báo nhu cầu thị trường và tối ưu hóa tuyến đường vận chuyển cho shipper. Từ lịch sử giao dịch, mô hình AI có thể dự báo nhu cầu từng loại nông sản theo mùa, giúp nông dân kế hoạch sản xuất hiệu quả hơn và giảm lãng phí."),
        ],
        "note": "💡 Nhấn mạnh: AI tư vấn theo mùa ĐÃ XONG. Hướng tương lai là Demand Forecasting & Route Optimization."
    },
    {
        "number": 18,
        "title": "Kết quả đạt được",
        "lines": [
            ("HUY", "Thưa thầy cô, về kết quả đạt được của nhóm em:\n\nVề mặt chức năng, nhóm em đã hoàn thành đầy đủ các tính năng đặt ra: quản lý sản phẩm và lô hàng, phân quyền 4 vai trò, ký giao dịch blockchain, truy xuất nguồn gốc theo thời gian thực, quản lý đơn hàng và tồn kho tự động, cùng hệ thống AI tư vấn tích hợp.\n\nVề mặt kỹ thuật, hệ thống đạt tốc độ phản hồi API dưới 100ms cho các tác vụ thông thường, và dưới 3 giây cho các yêu cầu cần đọc dữ liệu từ blockchain.\n\nĐặc biệt, hệ thống AI được thiết kế với ba lớp dự phòng — Gemini trực tiếp, OpenRouter và Mock Response — đảm bảo giao diện người dùng không bao giờ báo lỗi dù trong bất kỳ điều kiện mạng nào."),
        ],
        "note": ""
    },
    {
        "number": "19–22",
        "title": "Demo sản phẩm",
        "lines": [
            ("TIẾN", "Thưa thầy cô, bây giờ nhóm em xin phép trình bày demo sản phẩm thực tế. Nhóm em sẽ demo theo thứ tự luồng người dùng:\n\nĐầu tiên là trang chủ Cửa hàng — thầy cô có thể thấy Trợ lý Lão Nông AI đang đọc thời tiết thực tế và đưa ra tư vấn phù hợp.\n\nTiếp theo là trang Truy Xuất Nguồn Gốc — em nhập mã lô hàng 102 để tra cứu. Hệ thống sẽ kết nối thẳng với blockchain Sepolia và hiển thị toàn bộ hành trình của lô Sầu Riêng Ri6 này, bao gồm thời gian, vị trí và chữ ký của từng người trong chuỗi cung ứng.\n\nCuối cùng là trang Quản lý Kho — thầy cô có thể thấy dữ liệu tồn kho được đồng bộ tự động sau mỗi lần shipper xác nhận giao hàng thành công trên blockchain."),
        ],
        "note": "🔴 QUAN TRỌNG: Mở sẵn localhost:3000 trên trình duyệt. Demo lô hàng 102 (Sầu Riêng Ri6 - Cái Mơn, Bến Tre)."
    },
    {
        "number": 23,
        "title": "Kết luận",
        "lines": [
            ("HUY", "Thưa thầy cô, để tổng kết lại, VKU Market là một hệ thống ứng dụng thực tế giải quyết bài toán minh bạch chuỗi cung ứng nông sản bằng cách kết hợp ba công nghệ mạnh mẽ:\n\n• Blockchain — đảm bảo tính bất biến và không thể giả mạo của dữ liệu hành trình sản phẩm.\n• Web2 truyền thống — đảm bảo tốc độ và trải nghiệm mượt mà cho người dùng cuối không cần hiểu biết về crypto.\n• Trí tuệ nhân tạo — cầu nối thông minh biến dữ liệu kỹ thuật thành thông tin có ý nghĩa và giá trị thực tế cho người dùng.\n\nNhóm em tin rằng mô hình Web2.5 Hybrid này là hướng đi đúng đắn và thực tế nhất để đưa công nghệ Blockchain vào ứng dụng thương mại điện tử trong tương lai gần."),
        ],
        "note": ""
    },
    {
        "number": 24,
        "title": "Cảm ơn / Q&A",
        "lines": [
            ("TIẾN", "Thưa thầy cô, trên đây là toàn bộ nội dung thuyết trình của nhóm em. Em xin chân thành cảm ơn thầy cô đã lắng nghe và theo dõi. Nhóm em rất mong nhận được những nhận xét, góp ý quý báu từ hội đồng để nhóm em có thể hoàn thiện đề tài hơn nữa."),
            ("HUY", "Nhóm em xin kính mời thầy cô đặt câu hỏi. Chúng em rất sẵn sàng giải đáp."),
        ],
        "note": ""
    },
]

# ═══════════════════════════════════════════════════════════════
# RENDER SLIDES
# ═══════════════════════════════════════════════════════════════
SPEAKER_COLOR = {"HUY": TEAL, "TIẾN": RGBColor(0x6A, 0x1B, 0x9A)}
SPEAKER_ICON  = {"HUY": "🔵", "TIẾN": "🟣"}

for slide in slides:
    # Slide header band
    p = doc.add_paragraph()
    shade_paragraph(p, '146320')
    num = slide['number']
    add_run(p, f"  SLIDE {num}  ·  {slide['title']}", bold=True, color=WHITE, size=13, font='Calibri')

    # Speaker lines
    for speaker, text in slide['lines']:
        icon = SPEAKER_ICON.get(speaker, "")
        color = SPEAKER_COLOR.get(speaker, GRAY_DARK)
        # Speaker label
        lbl = doc.add_paragraph()
        lbl.paragraph_format.left_indent = Cm(0.5)
        lbl.paragraph_format.space_before = Pt(6)
        add_run(lbl, f"{icon} [{speaker}]:", bold=True, color=color, size=11)

        # Speech bubble content
        for line in text.split('\n'):
            sp = doc.add_paragraph()
            sp.paragraph_format.left_indent = Cm(1.5)
            sp.paragraph_format.space_before = Pt(1)
            sp.paragraph_format.space_after  = Pt(1)
            add_run(sp, line, size=11, color=GRAY_DARK)

    # Optional note box
    if slide.get('note'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        shade_paragraph(p, 'FFF9C4')
        add_run(p, f"  📝 GHI CHÚ: {slide['note']}", italic=True, color=RGBColor(0x5D, 0x40, 0x03), size=10)

    # ▶ CHUYỂN SLIDE
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    add_run(p, '  ▶ [CHUYỂN SLIDE]', bold=True, color=ORANGE, size=10)

    doc.add_paragraph()  # spacer

# ═══════════════════════════════════════════════════════════════
# Q&A APPENDIX
# ═══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
shade_paragraph(p, '0D47A1')
add_run(p, '  ❓  PHỤ LỤC: CÂU HỎI THƯỜNG GẶP & GỢI Ý TRẢ LỜI', bold=True, color=WHITE, size=14)
doc.add_paragraph()

qa_pairs = [
    ("Tại sao dùng Sepolia mà không phải Mainnet?",
     "Dạ thưa thầy/cô, trong giai đoạn nghiên cứu và phát triển của đồ án, nhóm em sử dụng Sepolia Testnet vì hai lý do chính.\n\nThứ nhất là chi phí: Ethereum Mainnet tốn phí gas thật, mỗi giao dịch có thể từ vài nghìn đến vài chục nghìn đồng. Trong môi trường học thuật, Sepolia cho phép nhóm em thử nghiệm không giới hạn mà không tốn tiền.\n\nThứ hai là an toàn: Hệ thống vẫn đang trong giai đoạn hoàn thiện. Sepolia đảm bảo nếu có lỗi, không ảnh hưởng đến tài sản thật của ai cả.\n\nKhi triển khai thực tế, nhóm em dự kiến chuyển sang Layer 2 như Arbitrum để có chi phí thấp hơn 95% so với Mainnet mà vẫn đảm bảo an toàn."),

    ("Nếu blockchain Sepolia bị tắt thì hệ thống có hoạt động không?",
     "Dạ thưa thầy/cô, đây là một câu hỏi rất hay. Hệ thống được thiết kế theo mô hình Graceful Degradation — tức là suy giảm một cách có kiểm soát.\n\nNếu blockchain không khả dụng, các tính năng thương mại điện tử cốt lõi — mua sắm, quản lý đơn hàng, tồn kho — vẫn hoạt động bình thường vì chúng dựa vào MySQL Off-chain.\n\nChỉ có tính năng Truy Xuất Nguồn Gốc là bị ảnh hưởng, vì nó phụ thuộc vào dữ liệu on-chain. Khi đó, hệ thống sẽ hiển thị thông báo lỗi thân thiện thay vì crash toàn bộ ứng dụng."),

    ("AI của em dùng mô hình gì? Có tốn tiền không?",
     "Dạ thưa thầy/cô, nhóm em sử dụng Google Gemini 2.5 Flash — đây là mô hình ngôn ngữ lớn của Google, rất mạnh về tiếng Việt và xử lý ngữ cảnh nông nghiệp.\n\nVề chi phí, nhóm em sử dụng API key trong giới hạn miễn phí của Google AI Studio nên hoàn toàn không tốn tiền trong giai đoạn hiện tại.\n\nNhóm em còn xây dựng cơ chế dự phòng ba lớp: nếu Gemini không khả dụng do đường truyền, hệ thống tự động chuyển sang OpenRouter với mô hình dự phòng miễn phí khác. Nếu cả hai đều không hoạt động, hệ thống phục vụ một tập câu trả lời được soạn sẵn rất chất lượng để đảm bảo Demo không bao giờ bị gián đoạn."),

    ("Dữ liệu cá nhân người dùng có được bảo vệ không?",
     "Dạ thưa thầy/cô, nhóm em đã thiết kế với nguyên tắc Privacy by Design.\n\nThông tin nhạy cảm như họ tên, số điện thoại, địa chỉ của khách hàng chỉ được lưu trong MySQL Off-chain với mã hóa và không bao giờ được ghi lên blockchain công khai.\n\nTrên blockchain, nhóm em chỉ lưu những thông tin phi cá nhân: ID lô hàng, trạng thái vận chuyển, tọa độ điểm giao hàng tổng quát và địa chỉ ví — mà địa chỉ ví không liên kết trực tiếp đến thông tin cá nhân thật của shipper."),

    ("Hệ thống có thể scale lên được không nếu có nhiều người dùng?",
     "Dạ thưa thầy/cô, đây cũng là một câu hỏi rất thực tế.\n\nVề phía Web2 — MySQL và Next.js API — có thể scale theo chiều ngang bằng cách thêm server mà không cần thay đổi kiến trúc.\n\nVề phía blockchain — đây là điểm bottleneck chính của Ethereum Mainnet vì chỉ xử lý khoảng 15 giao dịch mỗi giây. Đó chính là lý do hướng phát triển tương lai của nhóm em là chuyển sang Layer 2 Arbitrum, có thể xử lý hàng nghìn giao dịch mỗi giây với chi phí rất thấp.\n\nTrong phạm vi đề tài nghiên cứu và thử nghiệm, hệ thống hiện tại đáp ứng tốt cho quy mô vừa và nhỏ."),
]

for i, (q, a) in enumerate(qa_pairs, 1):
    # Question
    p = doc.add_paragraph()
    shade_paragraph(p, 'E3F2FD')
    add_run(p, f"  ❓ CÂU HỎI {i}: {q}", bold=True, color=RGBColor(0x0D, 0x47, 0xA1), size=11)

    # Answer
    for line in a.split('\n'):
        ap = doc.add_paragraph()
        ap.paragraph_format.left_indent = Cm(1)
        add_run(ap, line, size=11, color=GRAY_DARK)

    doc.add_paragraph()

# ── Tips box ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
shade_paragraph(p, 'E8F5E9')
tips = [
    "Nói chậm, rõ ràng — KHÔNG đọc slide, nhìn thầy cô mà nói.",
    "Khi demo: mở sẵn localhost:3000 trong trình duyệt, nhập mã lô hàng 102.",
    "Khi có câu hỏi khó: 'Dạ câu hỏi của thầy/cô rất hay, để em trả lời từng phần nhé...'",
    "Luôn bắt đầu trả lời bằng 'Dạ thưa thầy/cô' để thể hiện sự tôn trọng.",
    "Thời gian mục tiêu: 12-15 phút thuyết trình + 5-10 phút Q&A.",
]
add_run(p, "  💡 LƯU Ý KHI THUYẾT TRÌNH:\n", bold=True, color=GREEN_DARK, size=12)
for tip in tips:
    add_run(p, f"  • {tip}\n", size=10, color=GRAY_DARK)

# Final footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(p, '146320')
add_run(p, '\n  🎉  CHÚC HAI BẠN THUYẾT TRÌNH THÀNH CÔNG RỰC RỠ!  🎉\n', bold=True, color=WHITE, size=14)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = r'd:\code\Nam_3_ki2\Blockchain\docs\kichban.docx'
doc.save(out_path)
print("Da tao thanh cong: " + out_path)
